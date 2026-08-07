"""Build LLM-readable context from chat attachments."""

from __future__ import annotations

from pathlib import Path

from sqlalchemy.orm import Session

from app.db.models import ChatAttachment
from app.services.attachments import (
    DOCX_CONTENT_TYPE,
    absolute_path,
    is_image_content_type,
    is_text_extractable,
)

MAX_CHARS_PER_FILE = 12_000
MAX_CHARS_TOTAL = 36_000


def _clip(text: str, n: int = MAX_CHARS_PER_FILE) -> str:
    t = (text or "").strip()
    if len(t) <= n:
        return t
    return t[: n - 1] + "…"


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf not installed") from exc
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(f"[page {i + 1}]\n{t.strip()}")
        if sum(len(p) for p in parts) >= MAX_CHARS_PER_FILE:
            break
    return "\n\n".join(parts).strip()


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx not installed") from exc
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
        if sum(len(p) for p in parts) >= MAX_CHARS_PER_FILE:
            break
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
            if sum(len(p) for p in parts) >= MAX_CHARS_PER_FILE:
                break
    return "\n".join(parts).strip()


def extract_attachment_text(row: ChatAttachment) -> tuple[str, str]:
    """Return (kind, body). kind is text|pdf|docx|image|binary|error."""
    try:
        path = absolute_path(row.storage_path)
    except Exception as exc:
        return "error", f"(could not resolve file: {exc})"
    if not path.is_file():
        return "error", "(file missing on disk)"

    ctype = (row.content_type or "").lower()
    name = (row.filename or "").lower()

    if is_image_content_type(ctype) or name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
        return (
            "image",
            (
                "(Image file - binary content is not pasted into the prompt. "
                "Filename/type/size are authoritative hints. "
                "If the user asks what “this” is, treat this image as the referent; "
                "do not invent an unrelated paper or document.)"
            ),
        )

    if ctype == "application/pdf" or name.endswith(".pdf"):
        try:
            text = _read_pdf_text(path)
        except Exception as exc:
            return "error", f"(PDF text extraction failed: {exc})"
        if not text.strip():
            return (
                "pdf",
                "(PDF opened but no extractable text - may be scanned/image-only.)",
            )
        return "pdf", _clip(text)

    if ctype == DOCX_CONTENT_TYPE or name.endswith(".docx"):
        try:
            text = _read_docx_text(path)
        except Exception as exc:
            return "error", f"(Word docx extraction failed: {exc})"
        if not text.strip():
            return "docx", "(docx opened but no extractable text.)"
        return "docx", _clip(text)

    if is_text_extractable(row.filename or "", ctype):
        try:
            return "text", _clip(_read_text_file(path))
        except Exception as exc:
            return "error", f"(text read failed: {exc})"

    return "binary", "(unsupported for inline extraction)"


def build_attachments_prompt_block(
    db: Session, *, message_id: int | None, tenant_id: int
) -> str:
    if not message_id:
        return ""
    rows = (
        db.query(ChatAttachment)
        .filter(
            ChatAttachment.message_id == message_id,
            ChatAttachment.tenant_id == tenant_id,
        )
        .order_by(ChatAttachment.id.asc())
        .all()
    )
    if not rows:
        return ""

    lines = [
        "ATTACHED FILES (use these as the primary source when the user says "
        '"this", "the file", "the document", or similar):',
    ]
    used = 0
    for row in rows:
        kind, body = extract_attachment_text(row)
        header = (
            f"--- file: {row.filename} | type={row.content_type} | "
            f"size={row.size_bytes}b | kind={kind} ---"
        )
        chunk = f"{header}\n{body}"
        if used + len(chunk) > MAX_CHARS_TOTAL:
            remain = MAX_CHARS_TOTAL - used
            if remain > 200:
                lines.append(chunk[: remain - 1] + "…")
            lines.append("(further attachment content truncated)")
            break
        lines.append(chunk)
        used += len(chunk)

    lines.append(
        "Instruction: Answer from the attached file content above. "
        "Do not substitute an unrelated document from earlier chat context."
    )
    return "\n\n".join(lines)
